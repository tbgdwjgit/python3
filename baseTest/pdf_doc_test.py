__author__ = 'Test-YLL'

# -*- coding:utf-8 -*-

import PyPDF2
import docx

'''
安装PyPDF2
'''

pdfFileObj = open('E:\\用户手册.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print(pdfReader.numPages)
pageObj = pdfReader.getPage(23)
print(pageObj.extractText())

'''
安装python-docx
'''

doc = docx.Document('空调使用管理制度.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

#doc.paragraphs[1].text
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
